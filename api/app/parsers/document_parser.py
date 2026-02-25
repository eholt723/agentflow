# api/app/parsers/document_parser.py
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "xlsx", "xls", "jpg", "jpeg", "png", "webp"}
IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp"}
MAX_TEXT_CHARS = 12_000  # truncation ceiling before LLM calls

_MEDIA_TYPES = {
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
}


@dataclass
class ParsedDocument:
    filename: str
    extension: str
    text: Optional[str] = None
    rows: List[Dict[str, Any]] = field(default_factory=list)
    row_count: int = 0
    parse_error: Optional[str] = None
    image_bytes: Optional[bytes] = None   # set for image uploads; vision LLM transcribes
    image_media_type: Optional[str] = None

    @property
    def truncated_text(self) -> str:
        """Text trimmed to MAX_TEXT_CHARS for LLM calls."""
        if not self.text:
            return ""
        if len(self.text) > MAX_TEXT_CHARS:
            return self.text[:MAX_TEXT_CHARS] + "\n[... truncated ...]"
        return self.text


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return _parse_pdf(filename, ext, content)
    elif ext == "docx":
        return _parse_docx(filename, ext, content)
    elif ext == "txt":
        return _parse_txt(filename, ext, content)
    elif ext in {"csv", "xlsx", "xls"}:
        return _parse_tabular(filename, ext, content)
    elif ext in IMAGE_EXTENSIONS:
        return _parse_image(filename, ext, content)
    else:
        return ParsedDocument(
            filename=filename, extension=ext,
            parse_error=f"Unsupported file type: .{ext}",
        )


def _parse_pdf(filename: str, ext: str, content: bytes) -> ParsedDocument:
    import pdfplumber
    try:
        text_parts: List[str] = []
        page_count = 0
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        text = "\n".join(text_parts).strip()
        if text:
            logger.info("pdf_parse filename=%s pages=%d chars=%d", filename, page_count, len(text))
            return ParsedDocument(filename=filename, extension=ext, text=text)
        # Text layer empty — try PyMuPDF (handles more encoding types and design-tool PDFs)
        logger.info("pdf_pdfplumber_empty filename=%s pages=%d — falling back to pymupdf", filename, page_count)
        return _parse_pdf_pymupdf(filename, ext, content)
    except Exception as e:
        logger.error("pdf_parse_failed filename=%s error=%s", filename, e)
        return ParsedDocument(filename=filename, extension=ext, parse_error=str(e))


def _parse_pdf_pymupdf(filename: str, ext: str, content: bytes) -> ParsedDocument:
    """
    Fallback PDF parser using PyMuPDF.
    1. Tries text extraction (handles design-tool and non-standard encodings).
    2. If still empty, renders the first page to PNG and routes through the vision pipeline.
    """
    import fitz  # PyMuPDF
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts: List[str] = []
        for page in doc:
            t = page.get_text()
            if t.strip():
                text_parts.append(t)
        text = "\n".join(text_parts).strip()
        if text:
            logger.info("pdf_pymupdf filename=%s pages=%d chars=%d", filename, len(doc), len(text))
            return ParsedDocument(filename=filename, extension=ext, text=text)
        # Truly image-based PDF — render first page to PNG for vision model
        logger.info("pdf_pymupdf_empty filename=%s — rendering page 1 to image for vision", filename)
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale for legibility
        image_bytes = pix.tobytes("png")
        return ParsedDocument(
            filename=filename, extension=ext,
            image_bytes=image_bytes, image_media_type="image/png",
        )
    except Exception as e:
        logger.error("pdf_pymupdf_failed filename=%s error=%s", filename, e)
        return ParsedDocument(filename=filename, extension=ext, parse_error=str(e))


def _parse_docx(filename: str, ext: str, content: bytes) -> ParsedDocument:
    from docx import Document
    try:
        doc = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        logger.info("docx_parse filename=%s chars=%d", filename, len(text))
        return ParsedDocument(filename=filename, extension=ext, text=text)
    except Exception as e:
        logger.error("docx_parse_failed filename=%s error=%s", filename, e)
        return ParsedDocument(filename=filename, extension=ext, parse_error=str(e))


def _parse_txt(filename: str, ext: str, content: bytes) -> ParsedDocument:
    try:
        text = content.decode("utf-8", errors="replace")
        logger.info("txt_parse filename=%s chars=%d", filename, len(text))
        return ParsedDocument(filename=filename, extension=ext, text=text)
    except Exception as e:
        return ParsedDocument(filename=filename, extension=ext, parse_error=str(e))


def _parse_image(filename: str, ext: str, content: bytes) -> ParsedDocument:
    """Store raw image bytes for vision LLM transcription in analyze_runner."""
    media_type = _MEDIA_TYPES.get(ext, "image/jpeg")
    logger.info("image_parse filename=%s media_type=%s bytes=%d", filename, media_type, len(content))
    return ParsedDocument(
        filename=filename, extension=ext,
        image_bytes=content, image_media_type=media_type,
    )


def _parse_tabular(filename: str, ext: str, content: bytes) -> ParsedDocument:
    import pandas as pd
    try:
        if ext == "csv":
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_excel(io.BytesIO(content))
        rows = df.fillna("").to_dict(orient="records")
        text = df.to_string(index=False)
        logger.info(
            "tabular_parse filename=%s rows=%d cols=%d", filename, len(rows), len(df.columns)
        )
        return ParsedDocument(
            filename=filename, extension=ext,
            text=text, rows=rows, row_count=len(rows),
        )
    except Exception as e:
        logger.error("tabular_parse_failed filename=%s error=%s", filename, e)
        return ParsedDocument(filename=filename, extension=ext, parse_error=str(e))
